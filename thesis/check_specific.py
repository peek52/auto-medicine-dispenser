"""Check specific known-problem areas in the thesis docx."""

import sys
from docx import Document

DOCX_PATH = r"C:\Users\peekz\OneDrive\Desktop\mmmooo\project68\เล่มจบ\เล่มจบ1.docx"

doc = Document(DOCX_PATH)
ps = doc.paragraphs


def show(label, indices):
    print(f"\n--- {label} ---")
    for i in indices:
        if 0 <= i < len(ps):
            t = ps[i].text.strip()
            print(f"[p{i}] {t[:200]}")


# 1. Abstract (Thai)
print("=" * 60)
print("CHECK 1: บทคัดย่อ TH (around p136-160)")
print("=" * 60)
for i in range(136, 162):
    t = ps[i].text.strip()
    if t:
        print(f"[p{i}] {t[:250]}")

# 2. English abstract (after Thai)
print("\n" + "=" * 60)
print("CHECK 2: Abstract EN (paragraphs after Thai abstract)")
print("=" * 60)
# Find "Abstract" keyword
for i in range(150, 200):
    t = ps[i].text.strip()
    if "Abstract" in t or "Currently" in t or "WiFiManager" in t or "13 pills" in t:
        print(f"[p{i}] {t[:300]}")

# 3. §1.3.2 packaging
print("\n" + "=" * 60)
print("CHECK 3: §1.3.2 (p206)")
print("=" * 60)
show("around §1.3.2", range(204, 212))

# 4. §2.4 protocols — check if §2.4.2 exists
print("\n" + "=" * 60)
print("CHECK 4: §2.4 protocols (p352-360) — is §2.4.2 missing?")
print("=" * 60)
for i in range(350, 365):
    t = ps[i].text.strip()
    if t:
        print(f"[p{i}] {t[:200]}")

# 5. TOC
print("\n" + "=" * 60)
print("CHECK 5: TOC (table 0)")
print("=" * 60)
toc = doc.tables[0]
for row in toc.rows:
    cells = [c.text.strip()[:80] for c in row.cells]
    print("  ", " | ".join(cells))
