"""
Inspect a .docx file: list paragraphs, tables, and images.
Quick tool to understand structure before editing.

Usage:
    python thesis/inspect_docx.py <path-to-docx>
"""

import sys
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = sys.argv[1] if len(sys.argv) > 1 else \
    r"C:\Users\peekz\OneDrive\Desktop\mmmooo\project68\เล่มจบ\เล่มจบ1.docx"


def main():
    doc = Document(DOCX_PATH)
    print(f"=== INSPECTING: {DOCX_PATH} ===\n")

    # Count things
    paragraphs = doc.paragraphs
    tables = doc.tables
    sections = doc.sections

    # Count images via inline shapes + anchored shapes
    inline_images = len(doc.inline_shapes)

    print(f"Sections:        {len(sections)}")
    print(f"Paragraphs:      {len(paragraphs)}")
    print(f"Tables:          {len(tables)}")
    print(f"Inline images:   {inline_images}")
    print()

    # Find headings (paragraph styles starting with "Heading" or starting with chapter/section markers)
    print("=== STRUCTURE (headings + section markers) ===")
    chapter_count = 0
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        style = p.style.name if p.style else ""
        is_heading = style.startswith("Heading") or style.startswith("Title")

        # Detect Thai chapter / section markers
        is_chapter = (
            text.startswith("บทที่")
            or text.startswith("บทคัดย่อ")
            or text.startswith("กิตติกรรม")
            or text.startswith("สารบัญ")
            or text.startswith("ภาคผนวก")
            or text.startswith("บรรณานุกรม")
        )

        # Detect numbered sections like "1.1", "2.3.2", "3.5.1"
        first_word = text.split()[0] if text.split() else ""
        is_section = (
            len(first_word) <= 8
            and first_word.replace(".", "").isdigit()
            and "." in first_word
        )

        if is_heading or is_chapter or is_section:
            preview = text[:80] + ("..." if len(text) > 80 else "")
            marker = "📘" if is_chapter else ("§" if is_section else "▸")
            print(f"  [p{i:4d}] {marker} {preview}")
            if is_chapter:
                chapter_count += 1

    print(f"\nChapters detected: {chapter_count}")

    # Tables summary
    print(f"\n=== TABLES ({len(tables)}) ===")
    for i, tbl in enumerate(tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows > 0 else 0
        # Try to grab first cell as table preview
        first_cell = ""
        if rows > 0 and cols > 0:
            first_cell = tbl.rows[0].cells[0].text.strip()[:50]
        print(f"  [t{i:3d}] {rows} rows × {cols} cols — first cell: \"{first_cell}\"")


if __name__ == "__main__":
    main()
